"""Local resource loader — reads files from a ``resources/`` directory.

Supported formats:

* ``.md``, ``.txt`` — read as plain UTF-8 text (stdlib)
* ``.pdf``          — text extraction via *pymupdf*; scanned pages use
                      *pytesseract* + *Pillow* OCR.  Gracefully skipped when
                      the packages are not installed.
* ``.docx``         — paragraph text via *python-docx*.  Gracefully skipped
                      when the package is not installed.
* ``.epub``         — chapter text via *ebooklib* + *BeautifulSoup*.  Gracefully
                      skipped when the packages are not installed.
* ``.mobi``         — converted via *mobi* library.  Gracefully skipped when
                      not installed.

Directory scanning is **recursive** — sub-folders are traversed.

Each file becomes a ``SourceDocument`` with ``source_type="local_resource"``,
``acquisition_method="file_read"``, and ``url="file://…"``.
"""

from __future__ import annotations

import hashlib
import logging
import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

from ..state import SourceDocument

logger = logging.getLogger(__name__)

# Lazy import guards for optional heavy dependencies.
_HAS_PYMUPDF = False
try:
    import pymupdf  # noqa: F401

    _HAS_PYMUPDF = True
except ImportError:  # pragma: no cover
    pass

_HAS_PYTESSERACT = False
try:
    import pytesseract  # noqa: F401
    from PIL import Image  # noqa: F401

    _HAS_PYTESSERACT = True
except ImportError:  # pragma: no cover
    pass

_HAS_DOCX = False
try:
    import docx  # noqa: F401

    _HAS_DOCX = True
except ImportError:  # pragma: no cover
    pass

_HAS_EBOOKLIB = False
try:
    import ebooklib  # noqa: F401
    from bs4 import BeautifulSoup  # noqa: F401

    _HAS_EBOOKLIB = True
except ImportError:  # pragma: no cover
    pass

_HAS_MOBI = False
try:
    import mobi  # noqa: F401

    _HAS_MOBI = True
except ImportError:  # pragma: no cover
    pass

_TEXT_EXTENSIONS = {".md", ".txt", ".text", ".rst", ".csv", ".tsv", ".log"}
_MAX_CONTENT_CHARS = 80_000  # generous raw limit; individual nodes truncate further


# ---------------------------------------------------------------------------
# Format-specific readers
# ---------------------------------------------------------------------------

def _read_text_file(path: Path) -> Optional[str]:
    """Read a plain text file."""
    try:
        return path.read_text(encoding="utf-8", errors="replace")
    except OSError as exc:
        logger.warning("resource_loader: cannot read %s — %s", path, exc)
        return None


def _read_pdf(path: Path) -> Optional[str]:
    """Extract text from a PDF, with OCR fallback for scanned pages."""
    if not _HAS_PYMUPDF:
        logger.info(
            "resource_loader: skipping PDF %s — pymupdf not installed.  "
            "Install with: pip install 'sensemaking-agent[resources]'",
            path,
        )
        return None

    import pymupdf  # noqa: WPS433 — guarded import

    pages: list[str] = []
    try:
        doc = pymupdf.open(str(path))
    except Exception as exc:
        logger.warning("resource_loader: cannot open PDF %s — %s", path, exc)
        return None

    try:
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text") or ""
            # If little text but page has images, try OCR.
            if len(text.strip()) < 50 and _HAS_PYTESSERACT:
                text = _ocr_page(page) or text
            if text.strip():
                pages.append(text.strip())
    finally:
        doc.close()

    return "\n\n".join(pages) if pages else None


def _ocr_page(page: object) -> Optional[str]:
    """Run OCR on all images embedded in a pymupdf page object."""
    import io

    import pymupdf  # noqa: WPS433
    from PIL import Image as PILImage

    texts: list[str] = []
    try:
        image_list = page.get_images(full=True)  # type: ignore[attr-defined]
        doc = page.parent  # type: ignore[attr-defined]
        for img_meta in image_list:
            xref = img_meta[0]
            pix = pymupdf.Pixmap(doc, xref)
            if pix.n > 4:
                pix = pymupdf.Pixmap(pymupdf.csRGB, pix)
            img_bytes = pix.tobytes("png")
            pil_image = PILImage.open(io.BytesIO(img_bytes))
            ocr_text = pytesseract.image_to_string(pil_image)
            if ocr_text and ocr_text.strip():
                texts.append(ocr_text.strip())
    except Exception as exc:  # pragma: no cover
        logger.debug("resource_loader: OCR failed on page — %s", exc)
    return "\n".join(texts) if texts else None


def _read_docx(path: Path) -> Optional[str]:
    """Extract paragraph text from a .docx file."""
    if not _HAS_DOCX:
        logger.info(
            "resource_loader: skipping DOCX %s — python-docx not installed.  "
            "Install with: pip install 'sensemaking-agent[resources]'",
            path,
        )
        return None

    import docx as _docx  # noqa: WPS433 — guarded import

    try:
        document = _docx.Document(str(path))
    except Exception as exc:
        logger.warning("resource_loader: cannot open DOCX %s — %s", path, exc)
        return None

    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs) if paragraphs else None


def _read_epub(path: Path) -> Optional[str]:
    """Extract chapter text from an EPUB ebook."""
    if not _HAS_EBOOKLIB:
        logger.info(
            "resource_loader: skipping EPUB %s — ebooklib/beautifulsoup4 not installed.  "
            "Install with: pip install 'sensemaking-agent[resources]'",
            path,
        )
        return None

    from ebooklib import epub as _epub  # noqa: WPS433
    from bs4 import BeautifulSoup as _BS  # noqa: WPS433

    try:
        book = _epub.read_epub(str(path), options={"ignore_ncx": True})
    except Exception as exc:
        logger.warning("resource_loader: cannot open EPUB %s — %s", path, exc)
        return None

    chapters: list[str] = []
    for item in book.get_items_of_type(9):  # ITEM_DOCUMENT
        html = item.get_content().decode("utf-8", errors="replace")
        text = _BS(html, "html.parser").get_text(separator="\n", strip=True)
        if text:
            chapters.append(text)
    return "\n\n".join(chapters) if chapters else None


def _read_mobi(path: Path) -> Optional[str]:
    """Extract text from a MOBI ebook by converting to a temp directory."""
    if not _HAS_MOBI:
        logger.info(
            "resource_loader: skipping MOBI %s — mobi library not installed.  "
            "Install with: pip install 'sensemaking-agent[resources]'",
            path,
        )
        return None

    import shutil

    import mobi as _mobi  # noqa: WPS433

    try:
        tmpdir, extracted_path = _mobi.extract(str(path))
    except Exception as exc:
        logger.warning("resource_loader: cannot extract MOBI %s — %s", path, exc)
        return None

    try:
        extracted = Path(extracted_path)
        if not extracted.exists():
            logger.warning("resource_loader: MOBI extraction produced no output — %s", path)
            return None

        # The mobi library extracts to HTML; read and strip tags.
        try:
            raw = extracted.read_text(encoding="utf-8", errors="replace")
        except OSError as exc:
            logger.warning("resource_loader: cannot read extracted MOBI — %s", exc)
            return None

        if _HAS_EBOOKLIB:  # reuse BS4 if available
            from bs4 import BeautifulSoup as _BS  # noqa: WPS433
            return _BS(raw, "html.parser").get_text(separator="\n", strip=True) or None

        # Minimal fallback: strip obvious HTML tags.
        import re
        text = re.sub(r"<[^>]+>", " ", raw)
        text = re.sub(r"\s+", " ", text).strip()
        return text or None
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def _document_id(path: Path) -> str:
    """Derive a stable document ID from file path and modification time."""
    stat = path.stat()
    key = f"{path.resolve()}::{stat.st_mtime_ns}"
    return "doc_res_" + hashlib.sha1(key.encode()).hexdigest()[:16]


def load_resources(resources_dir: str | Path) -> list[SourceDocument]:
    """Read all supported files from *resources_dir* and return ``SourceDocument`` list.

    The directory is scanned **recursively** — sub-folders are traversed.
    Unsupported extensions are silently skipped with a debug log.
    """
    base = Path(resources_dir)
    if not base.is_dir():
        logger.warning("resource_loader: directory does not exist — %s", base)
        return []

    results: list[SourceDocument] = []

    for entry in sorted(base.rglob("*")):
        if not entry.is_file():
            continue

        ext = entry.suffix.lower()
        content: Optional[str] = None

        if ext in _TEXT_EXTENSIONS:
            content = _read_text_file(entry)
        elif ext == ".pdf":
            content = _read_pdf(entry)
        elif ext == ".docx":
            content = _read_docx(entry)
        elif ext == ".epub":
            content = _read_epub(entry)
        elif ext == ".mobi":
            content = _read_mobi(entry)
        else:
            logger.debug("resource_loader: skipping unsupported file %s", entry.name)
            continue

        if not content or not content.strip():
            logger.debug("resource_loader: empty content from %s — skipping", entry.name)
            continue

        # Truncate excessively large content.
        if len(content) > _MAX_CONTENT_CHARS:
            content = content[:_MAX_CONTENT_CHARS]
            logger.warning(
                "resource_loader: truncated %s to %d chars.", entry.name, _MAX_CONTENT_CHARS
            )

        doc = SourceDocument(
            document_id=_document_id(entry),
            url=entry.resolve().as_uri(),
            title=entry.name,
            content=content,
            source_type="local_resource",
            query="",  # not query-driven
            acquisition_method="file_read",
            metadata={"original_path": str(entry.resolve()), "size_bytes": entry.stat().st_size},
        )
        results.append(doc)

    logger.info("resource_loader: loaded %d documents from %s.", len(results), base)
    return results
